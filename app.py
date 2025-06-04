import os
import streamlit as st
from datetime import datetime
from docx import Document
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from dotenv import load_dotenv
import json
import logging

# Carrega variáveis de ambiente
@st.cache_data
def carregar_configuracoes():
    load_dotenv()
    return {
        "EMAIL_USER": os.getenv("EMAIL_USER"),
        "EMAIL_PASS": os.getenv("EMAIL_PASS"),
        "EMAIL_SMTP": os.getenv("EMAIL_SMTP"),
        "EMAIL_PORT": int(os.getenv("EMAIL_PORT", 587))
    }

# Configuração da página
st.set_page_config(page_title="Formulário de Pré-Avaliação Neuropsicológica", layout="wide")

# Título e instruções iniciais
st.title("Formulário de Pré-Avaliação Neuropsicológica")

st.markdown("""
Este formulário tem como objetivo coletar informações iniciais para sua avaliação neuropsicológica.  
Os dados serão mantidos sob sigilo e utilizados apenas para fins clínicos, conforme o Código de Ética Profissional do Psicólogo (CFP).
""")

# Checkbox de consentimento
consentimento = st.checkbox("Declaro que li e concordo com o uso das informações acima para fins de avaliação psicológica.")

# Exibe formulário somente se o consentimento for marcado
if consentimento:
    with st.form(key="form_anamnese"):
        # == 1. Dados Pessoais ==
        st.header("1. Dados Pessoais")
        nome = st.text_input("Nome completo:", max_chars=100)
        email = st.text_input("E-mail para contato:")
        telefone = st.text_input("Telefone:")
        data_avaliacao = st.date_input(
            "Data da Avaliação:",
            datetime.today(),
            format="DD/MM/YYYY"
        )
        data_nasc = st.date_input(
            "Data de Nascimento:",
            min_value=datetime(1900, 1, 1),
            max_value=datetime.today(),
            format="DD/MM/YYYY"
        )
        idade = st.number_input("Idade:", min_value=0, max_value=120, step=1)
        sexo = st.selectbox("Sexo:", ["Masculino", "Feminino", "Outro"])
        estado_nasc = st.text_input("Estado de nascimento:")
        cidade_nasc = st.text_input("Cidade de nascimento:")
        endereco = st.text_input("Endereço completo:")
        mao_escrita = st.radio("Mão que usa para escrever:", ["Direita", "Esquerda", "Ambidestro"])
        idiomas = st.multiselect(
            "Fala outro(s) idioma(s)?",
            ["Não", "Inglês", "Espanhol", "Francês", "Outro"]
        )
        encaminhamento = st.text_input("Se você foi encaminhado(a) por algum profissional de saúde, informe aqui:")

        st.markdown("---")
        # == 2. Queixas Principais ==
        st.header("2. Queixas Principais")
        queixas = st.text_area("Descreva brevemente o motivo para a avaliação:", height=80)

        st.markdown("---")
        # == 3. Sintomas Cognitivos ==
        st.header("3. Sintomas Cognitivos")
        cog_concentracao = st.radio("Apresenta dificuldade de concentração?", ["Sim", "Não"])
        cog_esquecimento = st.radio("Apresenta esquecimento frequente de compromissos ou informações recentes?", ["Sim", "Não"])
        cog_raciocinio = st.radio("Apresenta lentidão no raciocínio?", ["Sim", "Não"])
        cog_perda_objetos = st.radio("Apresenta perda de objetos com frequência?", ["Sim", "Não"])
        cog_repeticao = st.radio("Apresenta repetição de perguntas ou frases?", ["Sim", "Não"])
        cog_foco = st.radio("Apresenta dificuldade em manter o foco durante conversas?", ["Sim", "Não"])
        cog_desorientacao = st.radio("Apresenta sensação de desorientação (tempo, espaço, pessoas)?", ["Sim", "Não"])
        cog_problemas = st.radio("Apresenta dificuldade para resolver problemas cotidianos?", ["Sim", "Não"])
        cog_lembretes = st.radio("Apresenta necessidade constante de listas ou lembretes?", ["Sim", "Não"])
        cog_cansaco = st.radio("Apresenta sensação de cansaço mental excessivo após esforço intelectual?", ["Sim", "Não"])
        cog_palavras = st.radio("Apresenta troca ou inversão de palavras ao falar ou escrever?", ["Sim", "Não"])
        cog_anomia = st.radio("Apresenta dificuldade para encontrar palavras durante a fala (anomia)?", ["Sim", "Não"])

        st.markdown("---")
        # == 4. Histórico Médico: Situação Atual e Passada ==
        st.header("4. Histórico Médico: Situação Atual e Passada")
        condicoes_medicas = st.multiselect(
            "Marque as condições médicas que se aplicam:",
            [
                "Hipertensão arterial",
                "Diabetes",
                "Doenças cardíacas",
                "Doenças neurológicas",
                "Doenças psiquiátricas",
                "Câncer",
                "Outra(s) condição(ões) relevante(s)"
            ]
        )

        if "Câncer" in condicoes_medicas:
            cancer_info = st.text_input("Especifique o tipo de câncer e a data do diagnóstico:")

        if "Doenças psiquiátricas" in condicoes_medicas:
            psiquiatria_info = st.text_input("Especifique o diagnóstico psiquiátrico:")

        if "Outra(s) condição(ões) relevante(s)" in condicoes_medicas:
            outras_condicoes = st.text_input("Descreva outras condições médicas relevantes:")

        st.subheader("Uso de Medicações")
        usa_medicacao = st.radio("Faz uso contínuo ou recente de medicações?", ["Não", "Sim"])
        medicacoes = st.text_area(
            "Nome do(s) medicamento(s) – Dosagem – Motivo – Por quem foi prescrito:",
            height=80  # mínimo 68px
        )

        historico_medico = st.text_area("Descreva seu histórico médico atual e passado:", height=80)
        historico_familiar = st.text_area("Descreva histórico médico familiar:", height=80)

        st.markdown("---")
        # == 5. Aspectos do Desenvolvimento Infantil ==
        st.header("5. Aspectos do Desenvolvimento Infantil")
        st.markdown("""
**Como foi o desenvolvimento na infância? Indique se você teve alguma dificuldade em algum dos aspectos a seguir:**

- **Desenvolvimento motor grosso:** Sentar-se sem apoio, engatinhar, andar sozinho, correr, pular e subir escadas.  
- **Desenvolvimento motor fino:** Agilidade para segurar objetos pequenos com precisão, usar talheres, lápis ou tesoura, vestir-se e despir-se de forma autônoma.  
- **Desenvolvimento da fala e linguagem:** Emissão de sons na idade esperada, formação de palavras e frases, compreensão de comandos e clareza na articulação.  
- **Desenvolvimento cognitivo:** Resolver problemas simples, reconhecimento de formas, cores e números, memória e raciocínio.  
- **Desenvolvimento emocional:** Expressão adequada de sentimentos (alegria, frustração, medo), controle emocional em situações desafiadoras, vinculação afetiva com cuidadores.  
- **Desenvolvimento social:** Interação com adultos e outras crianças, partilha de brinquedos, participação em atividades em grupo, seguir regras simples.  
- **Autonomia:** Higiene pessoal (lavar as mãos, escovar os dentes), controle de esfíncteres, alimentação independente.  
- **Aquisição de hábitos e rotinas:** Sono regular, alimentação equilibrada, adaptação ao ambiente escolar.
""")
        desenvolvimento_infantil = st.text_area(
            "Relate detalhadamente o desenvolvimento infantil conforme descrito acima:",
            height=120
        )

        st.markdown("---")
        # == 6. Aspectos do Desenvolvimento Escolar ==
        st.header("6. Aspectos do Desenvolvimento Escolar")
        historico_escolar = st.text_area(
            "Descreva como foi a escolarização (repetições, dificuldades, apoio pedagógico):",
            height=100
        )

        st.markdown("---")
        # == 7. Aspectos Emocionais ==
        st.header("7. Aspectos Emocionais")
        emocional_sono = st.radio("Alterações de sono?", ["Sim", "Não"])
        emocional_apetite = st.radio("Alterações de apetite?", ["Sim", "Não"])
        emocional_humor = st.radio("Oscilações de humor ou tristeza frequente?", ["Sim", "Não"])
        emocional_estresse = st.radio("Nível de estresse percebido:", ["Baixo", "Moderado", "Alto"])

        st.markdown("---")
        # == 8. Uso de Neurotecnologias (Opcional) ==
        st.header("8. Uso de Neurotecnologias")
        uso_neuro = st.text_area(
            "Informe se já utiliza Neurofeedback, tDCS, Hipnose, BrainTap, Muse ou outros:",
            height=80
        )

        st.markdown("---")
        # == 9. Observações Finais ==
        st.header("9. Observações Finais")
        observacoes = st.text_area("Informações adicionais relevantes:", height=80)

        submit_button = st.form_submit_button(label="Enviar Avaliação")

        if submit_button:
            if st.button("Confirmar envio dos dados?"):
                with st.spinner('Gerando documento...'):
                    # Sanitize e define nome do arquivo Word
                    nome_sanitizado = nome.strip().replace(" ", "_")
                    filename = f"avaliacao_{nome_sanitizado}.docx"
                    doc = Document()

                    # == Cabeçalho com Dados Pessoais ==
                    doc.add_heading(f"Pré-Avaliação Neuropsicológica: {nome}", level=1)
                    doc.add_paragraph(f"Data da Avaliação: {data_avaliacao.strftime('%d/%m/%Y')}")
                    doc.add_paragraph(f"E-mail: {email if email else 'Não informado'}")
                    doc.add_paragraph(f"Telefone: {telefone if telefone else 'Não informado'}")
                    doc.add_paragraph(f"Nascimento: {data_nasc.strftime('%d/%m/%Y')}  (Idade: {idade})")
                    doc.add_paragraph(f"Sexo: {sexo}")
                    doc.add_paragraph(f"Cidade/Estado de nascimento: {cidade_nasc}/{estado_nasc}")
                    doc.add_paragraph(f"Endereço: {endereco if endereco else 'Não informado'}")
                    doc.add_paragraph(f"Mão dominante: {mao_escrita}")
                    idi_list = [i for i in idiomas if i != "Não"]
                    doc.add_paragraph(f"Idiomas: {', '.join(idi_list) if idi_list else 'Nenhum'}")
                    doc.add_paragraph(f"Encaminhado por: {encaminhamento if encaminhamento else 'Nenhum'}")

                    # == Seção 2: Queixas Principais ==
                    doc.add_page_break()
                    doc.add_heading("2. Queixas Principais", level=2)
                    doc.add_paragraph(queixas if queixas else "Nenhuma")

                    # == Seção 3: Sintomas Cognitivos ==
                    doc.add_page_break()
                    doc.add_heading("3. Sintomas Cognitivos", level=2)
                    doc.add_paragraph(f"Dificuldade de concentração: {cog_concentracao}")
                    doc.add_paragraph(f"Esquecimento frequente: {cog_esquecimento}")
                    doc.add_paragraph(f"Lentidão no raciocínio: {cog_raciocinio}")
                    doc.add_paragraph(f"Perda de objetos: {cog_perda_objetos}")
                    doc.add_paragraph(f"Repetição de perguntas/frases: {cog_repeticao}")
                    doc.add_paragraph(f"Dificuldade de foco em conversas: {cog_foco}")
                    doc.add_paragraph(f"Sensação de desorientação: {cog_desorientacao}")
                    doc.add_paragraph(f"Dificuldade para resolver problemas cotidianos: {cog_problemas}")
                    doc.add_paragraph(f"Necessidade de listas/lembretes: {cog_lembretes}")
                    doc.add_paragraph(f"Cansaço mental excessivo: {cog_cansaco}")
                    doc.add_paragraph(f"Troca/inversão de palavras: {cog_palavras}")
                    doc.add_paragraph(f"Dificuldade de encontrar palavras (anomia): {cog_anomia}")

                    # == Seção 4: Histórico Médico ==
                    doc.add_page_break()
                    doc.add_heading("4. Histórico Médico: Situação Atual e Passada", level=2)
                    doc.add_paragraph(f"Condições Médicas selecionadas: {', '.join(condicoes_medicas) if condicoes_medicas else 'Nenhuma'}")
                    if "Câncer" in condicoes_medicas:
                        doc.add_paragraph(f"   • Tipo de câncer e data do diagnóstico: {cancer_info if cancer_info else 'Não informado'}")
                    if "Doenças psiquiátricas" in condicoes_medicas:
                        doc.add_paragraph(f"   • Diagnóstico psiquiátrico: {psiquiatria_info if psiquiatria_info else 'Não informado'}")
                    if "Outra(s) condição(ões) relevante(s)" in condicoes_medicas:
                        doc.add_paragraph(f"   • Outras condições médicas: {outras_condicoes if outras_condicoes else 'Não informado'}")

                    doc.add_heading("Uso de Medicações", level=3)
                    doc.add_paragraph(f"Faz uso contínuo ou recente de medicações: {usa_medicacao}")
                    doc.add_paragraph(f"   Medicamentos informados: {medicacoes if medicacoes else 'Nenhum'}")
                    doc.add_paragraph(f"Histórico Médico Pessoal: {historico_medico if historico_medico else 'Não descrito'}")
                    doc.add_paragraph(f"Histórico Médico Familiar: {historico_familiar if historico_familiar else 'Não descrito'}")

                    # == Seção 5: Aspectos do Desenvolvimento Infantil ==
                    doc.add_page_break()
                    doc.add_heading("5. Aspectos do Desenvolvimento Infantil", level=2)
                    doc.add_paragraph(desenvolvimento_infantil if desenvolvimento_infantil else "Não descrito")

                    # == Seção 6: Aspectos do Desenvolvimento Escolar ==
                    doc.add_page_break()
                    doc.add_heading("6. Aspectos do Desenvolvimento Escolar", level=2)
                    doc.add_paragraph(historico_escolar if historico_escolar else "Não descrito")

                    # == Seção 7: Aspectos Emocionais ==
                    doc.add_page_break()
                    doc.add_heading("7. Aspectos Emocionais", level=2)
                    doc.add_paragraph(f"Alterações de sono: {emocional_sono}")
                    doc.add_paragraph(f"Alterações de apetite: {emocional_apetite}")
                    doc.add_paragraph(f"Oscilações de humor/tristeza: {emocional_humor}")
                    doc.add_paragraph(f"Nível de estresse percebido: {emocional_estresse}")

                    # == Seção 8: Uso de Neurotecnologias ==
                    doc.add_page_break()
                    doc.add_heading("8. Uso de Neurotecnologias", level=2)
                    doc.add_paragraph(uso_neuro if uso_neuro else "Nenhum uso informado")

                    # == Seção 9: Observações Finais ==
                    doc.add_page_break()
                    doc.add_heading("9. Observações Finais", level=2)
                    doc.add_paragraph(observacoes if observacoes else "Nenhuma observação adicional")

                    # Salva o arquivo
                    doc.save(filename)
                    st.success(f"Arquivo `{filename}` gerado com sucesso.")

                with st.spinner('Enviando email...'):
                    # Envio por e-mail
                    try:
                        msg = MIMEMultipart()
                        msg["From"] = EMAIL_USER
                        msg["To"] = email or EMAIL_USER
                        msg["Subject"] = f"Avaliação de {nome} – {data_avaliacao.strftime('%d/%m/%Y')}"

                        with open(filename, "rb") as attachment:
                            part = MIMEBase(
                                "application",
                                "vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                            part.set_payload(attachment.read())
                            encoders.encode_base64(part)
                            part.add_header("Content-Disposition", f'attachment; filename="{filename}"')
                            msg.attach(part)

                        server = smtplib.SMTP(EMAIL_SMTP, EMAIL_PORT)
                        server.ehlo()            # Handshake inicial antes do STARTTLS
                        server.starttls()
                        server.ehlo()            # Handshake após TLS
                        server.login(EMAIL_USER, EMAIL_PASS)
                        server.sendmail(EMAIL_USER, [email or EMAIL_USER], msg.as_string())
                        server.quit()

                        st.success("E-mail enviado com sucesso.")

                        # Após enviar o email, limpar o arquivo temporário
                        if os.path.exists(filename):
                            os.remove(filename)
                    except Exception as e:
                        st.error(f"Erro ao enviar e-mail: {e}")

def salvar_backup(dados):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_file = f"backup_{timestamp}.json"
    with open(backup_file, 'w') as f:
        json.dump(dados, f)

logging.basicConfig(
    filename='form_log.txt',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
